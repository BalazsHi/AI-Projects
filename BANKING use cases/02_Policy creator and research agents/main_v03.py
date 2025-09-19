import os
import shutil
import re
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langgraph.graph import StateGraph, END
from typing import List, TypedDict, Dict
from docx import Document
from serpapi import GoogleSearch # Ensure serpapi is installed: pip install google-search-results
from pypdf import PdfReader # Import PdfReader for reading PDFs

# --- Configuration ---
# Make sure to set your GOOGLE_API_KEY and SERPAPI_API_KEY as environment variables
if "GOOGLE_API_KEY" not in os.environ:
    raise ValueError("GOOGLE_API_KEY environment variable not set.")

if "SERPAPI_API_KEY" not in os.environ:
    raise ValueError("SERPAPI_API_KEY environment variable not set.")

# --- Document Handling ---
def save_docx(filepath, content):
    doc = Document()
    doc.add_heading('Internal Credit Risk Modelling Policy', 0)
    # Split content into paragraphs for better document structure
    for paragraph_text in content.split('\n\n'):
        doc.add_paragraph(paragraph_text)
    doc.save(filepath)

def read_docx(filepath):
    doc = Document(filepath)
    return "\n".join([p.text for p in doc.paragraphs])

def read_pdf(filepath):
    text_content = ""
    try:
        reader = PdfReader(filepath)
        for page in reader.pages:
            text_content += page.extract_text() + "\n"
    except Exception as e:
        print(f"Error reading PDF {filepath}: {e}")
    return text_content

# --- Agentic Workflow State ---
class WorkflowState(TypedDict):
    initial_documents_content: List[str] # Stores content of user-provided PDFs
    regulations: List[Dict[str, str]]
    drafts: List[str]
    final_policy: str

# --- Agents ---

# New: 0. Load Initial Documents Agent
def load_initial_documents_agent(state: WorkflowState):
    print("\n--- LOADING USER-PROVIDED DOCUMENTS ---") # Added newline for better visibility
    input_folder = "input_documents"
    initial_docs_content_list = [] # Renamed local variable for clarity

    if not os.path.exists(input_folder):
        print(f"ERROR: The '{input_folder}' folder was not found. Please create it and place your regulatory documents inside.")
        state["initial_documents_content"] = [] # Ensure state is explicitly empty if folder missing
        return state

    files_in_folder = [f for f in os.listdir(input_folder) if os.path.isfile(os.path.join(input_folder, f))]
    if not files_in_folder:
        print(f"INFO: The '{input_folder}' folder is empty. No user-provided documents to load.")
        state["initial_documents_content"] = []
        return state

    print(f"Found {len(files_in_folder)} items in '{input_folder}'. Processing...")

    for filename in files_in_folder:
        filepath = os.path.join(input_folder, filename)
        if filename.lower().endswith(".pdf"):
            print(f"Attempting to read PDF: {filename}")
            content = read_pdf(filepath)
            if content:
                initial_docs_content_list.append(content)
                print(f"Successfully extracted {len(content)} characters from {filename}.")
            else:
                print(f"Warning: No content could be extracted from PDF: {filename}. It might be empty, corrupted, or password-protected.")
        elif filename.lower().endswith(".txt"):
            print(f"Attempting to read TXT: {filename}")
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
                if content:
                    initial_docs_content_list.append(content)
                    print(f"Successfully extracted {len(content)} characters from {filename}.")
                else:
                    print(f"Warning: TXT file {filename} is empty.")
            except Exception as e:
                print(f"Error reading TXT file {filename}: {e}")
        elif filename.lower().endswith(".docx"):
            print(f"Attempting to read DOCX: {filename}")
            try:
                content = read_docx(filepath)
                if content:
                    initial_docs_content_list.append(content)
                    print(f"Successfully extracted {len(content)} characters from {filename}.")
                else:
                    print(f"Warning: DOCX file {filename} is empty or no readable paragraphs.")
            except Exception as e:
                print(f"Error reading DOCX file {filename}: {e}")
        else:
            print(f"Skipping unsupported file type: {filename}")
                
    state["initial_documents_content"] = initial_docs_content_list
    print(f"--- Finished loading user-provided documents. Total loaded: {len(state['initial_documents_content'])} documents. ---")
    return state


# 1. Research Agent (No changes needed here for reading initial docs, but it runs after them)
def research_agent(state: WorkflowState):
    print("--- RESEARCH AGENT ---")
    llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash")

    query_gen_prompt = ChatPromptTemplate.from_template(
        "You are an expert in EU banking regulation. Generate 3-5 highly specific Google search queries "
        "to find official European Banking Authority (EBA) regulations, guidelines, and regulatory technical standards (RTS) "
        "related to Internal Ratings-Based (IRB) model development for credit risk, specifically for residential mortgage portfolios. "
        "Focus on documents relevant to a retail bank operating in the European Union. "
        "Prioritize documents directly from EBA or official EU legal publications (e.g., eur-lex.europa.eu). "
        "Provide only the search queries, one per line, without any introductory or concluding remarks."
    )
    query_gen_chain = query_gen_prompt | llm | StrOutputParser()
    
    search_queries_str = query_gen_chain.invoke({})
    search_queries = [q.strip() for q in search_queries_str.split('\n') if q.strip()]
    
    print(f"Generated search queries: {search_queries}")

    found_regulations = []
    
    for query in search_queries:
        try:
            params = {
                "engine": "google",
                "q": query,
                "api_key": os.environ["SERPAPI_API_KEY"],
                "hl": "en"
            }
            search = GoogleSearch(params)
            results = search.get_dict()

            for result in results.get("organic_results", []):
                title = result.get("title")
                snippet = result.get("snippet")
                link = result.get("link")

                if link and ("eba.europa.eu" in link or "eur-lex.europa.eu" in link): # Removed dnb.nl as per original prompt focus
                    found_regulations.append({
                        "title": title,
                        "summary": snippet,
                        "link": link,
                        "content": snippet 
                    })
                    if len(found_regulations) >= 5:
                        break
            if len(found_regulations) >= 5:
                break
        except Exception as e:
            print(f"Error during search for query '{query}': {e}")

   
    if not os.path.exists("research_results"):
        os.makedirs("research_results")

    for i, reg in enumerate(found_regulations):
        filename = f"research_results/regulation_{i+1}.txt"
        with open(filename, "w", encoding="utf-8") as f:
            f.write(f"Title: {reg['title']}\n")
            f.write(f"Link: {reg['link']}\n")
            f.write(f"Summary/Content Snippet:\n{reg['content']}\n")
        reg["filepath"] = filename 
        
    state["regulations"] = found_regulations
    return state

# 2. Reader/Drafter Agent (Modified to accept initial_documents_content)
def reader_drafter_agent(regulation: Dict[str, str], initial_documents_content: List[str]):
    print(f"--- READER/DRAFTER AGENT for {regulation['title']} ---")
    llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash")

    # Prepare initial documents for the prompt
    initial_docs_text = "\n\n--- User Provided Document ---\n".join(initial_documents_content) if initial_documents_content else "No additional user-provided documents."

    prompt = ChatPromptTemplate.from_template(
        "You are an expert in EU banking regulation and credit risk. Your task is to read the provided "
        "regulatory text or summary related to IRB model development for residential mortgages. "
        "**Crucially, also consider the following user-provided regulatory documents and incorporate their requirements and insights.** " # New instruction
        "Based on these, draft a detailed chapter or section for an internal policy on IRB model development "
        "for a retail bank operating in the European Union, focusing on residential mortgage portfolios. "
        "The draft should extract key requirements, methodologies, data standards, "
        "risk differentiation, risk quantification, Appropriate Adjustment, Margin of Conservtism, model performance assessment, "
        "and any other relevant factor for credit risk model development. Ensure the language is formal, precise, and directly applicable to an "
        "internal policy document. Include specific references to the regulation's title or key concepts "
        "where appropriate.Start with the User Provided Regulatory Documents. \n\n"
        "--- User Provided Regulatory Documents ---\n{initial_docs_text}"
        "Regulatory Document Details:\nTitle: {title}\nLink: {link}\n\n"
        "Regulatory Text/Content:\n{text}\n\n"
         # New input
    )

    chain = prompt | llm | StrOutputParser()
    
    text_content = ""
    if "filepath" in regulation and os.path.exists(regulation["filepath"]):
        with open(regulation["filepath"], "r", encoding="utf-8") as f:
            text_content = f.read()
    else:
        text_content = regulation.get("content", regulation.get("summary", "No content available."))

    draft = chain.invoke({"title": regulation['title'], "link": regulation['link'], "text": text_content, "initial_docs_text": initial_docs_text})

    if not os.path.exists("output_documents"):
        os.makedirs("output_documents")
    
    sanitized_title = re.sub(r'[^\w\-_\. ]', '', regulation['title'])
    filepath = f"output_documents/draft_{sanitized_title}.docx"
    save_docx(filepath, draft)

    return draft

def parallel_drafting(state: WorkflowState):
    print("--- PARALLEL DRAFTING ---")
    # Pass initial_documents_content to each reader_drafter_agent call
    drafts = [reader_drafter_agent(reg, state["initial_documents_content"]) for reg in state["regulations"]]
    state["drafts"] = drafts
    return state

# 3. Synthesizing Agent (Modified to accept initial_documents_content)
def synthesizing_agent(state: WorkflowState):
    print("--- SYNTHESIZING AGENT ---")
    llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash")

    # Prepare initial documents for the prompt
    initial_docs_text = "\n\n--- User Provided Document ---\n".join(state["initial_documents_content"]) if state["initial_documents_content"] else "No additional user-provided documents."

    prompt = ChatPromptTemplate.from_template(
        "You are a senior credit risk model developer and experienced policy writer for a retail bank in the European Union. "
        "Your task is to synthesize the following draft chapters into a single, coherent, "
        "and well-structured internal policy document on IRB model development for residential mortgages. "
        "**Ensure all requirements and insights from the provided drafts AND the user-provided regulatory documents are fully integrated.** " # New instruction
        "The final document should have a logical structure with an introduction, main sections "
        "(e.g., General Principles, PD Model Requirements, LGD Model Requirements, Data Standards, "
        "Appropriate Adjustment and Margin of Conservatism), "
        "a table of contents, clear headings, and be comprehensive. "
        "Eliminate any duplications, resolve inconsistencies, and ensure the language is clear, "
        "professional, and transparent. The policy should be at least 10 pages long to ensure sufficient detail, "
        "reflecting the stringent regulatory environment in the EU/Netherlands for IRB models.\n\n"
        "Drafts to synthesize:\n{drafts}\n\n"
        "--- User Provided Regulatory Documents for Consideration ---\n{initial_docs_text}" # New input
    )

    chain = prompt | llm | StrOutputParser()
    final_policy = chain.invoke({"drafts": "\n\n---\n\n".join(state["drafts"]), "initial_docs_text": initial_docs_text})
    
    state["final_policy"] = final_policy
    save_docx("output_documents/Final_Internal_Policy.docx", final_policy)
    return state

# --- LangGraph Workflow Definition ---
graph = StateGraph(WorkflowState)

# Add the new node for loading initial documents
graph.add_node("load_initial_documents", load_initial_documents_agent)
graph.add_node("research", research_agent)
graph.add_node("parallel_drafting", parallel_drafting)
graph.add_node("synthesize", synthesizing_agent)

# Set the new entry point and add a new edge
graph.set_entry_point("load_initial_documents")
graph.add_edge("load_initial_documents", "research")
graph.add_edge("research", "parallel_drafting")
graph.add_edge("parallel_drafting", "synthesize")
graph.add_edge("synthesize", END)

app = graph.compile()

# --- Running the Workflow ---
if __name__ == "__main__":
    # Clean up previous runs
    if os.path.exists("output_documents"):
        shutil.rmtree("output_documents")
    if os.path.exists("research_results"):
        shutil.rmtree("research_results")

    # Ensure 'input_documents' folder exists and prompt user if needed
    input_folder_path = "input_documents"
    if not os.path.exists(input_folder_path):
        os.makedirs(input_folder_path)
        print(f"\nINFO: Created '{input_folder_path}' folder.")
        print("Please place your PDF (or TXT/DOCX) regulatory documents into this folder.")
        print("The script will continue, but no user-provided documents will be processed until you add them.")
        # Optionally, you could add a pause here if you want to wait for user action
        # input("Press Enter after placing your documents in 'input_documents'...")
    else:
        print(f"\nINFO: '{input_folder_path}' folder already exists. Checking for documents...")

    # Initial state for the workflow
    initial_state = {"initial_documents_content": [], "regulations": [], "drafts": [], "final_policy": ""}
    
    print("\n--- Starting AI Agent Workflow ---")
    final_state = app.invoke(initial_state)

    print("\n--- FINAL POLICY ---")
    # Print a snippet of the final policy, as it can be very long
    if final_state["final_policy"]:
        print(f"Final policy generated (first 500 chars):\n{final_state['final_policy'][:500]}...")
    else:
        print("No final policy was generated.")

    print("\nWorkflow complete. Check the 'output_documents' folder for the final policy.")
    
    # --- Debugging: Print the content of initial_documents_content from the final state ---
    print(f"\n--- Debug Info: initial_documents_content in final_state ---")
    if final_state["initial_documents_content"]:
        print(f"Found {len(final_state['initial_documents_content'])} user-provided documents in final state.")
        for i, doc_content in enumerate(final_state["initial_documents_content"]):
            print(f"  Doc {i+1} (length {len(doc_content)} characters):")
            print(f"    Content snippet: '{doc_content[:min(200, len(doc_content))]}...'") # Print first 200 chars
    else:
        print("No user-provided documents found in the final state's 'initial_documents_content'. This means either:")
        print("  1. The 'input_documents' folder was empty or did not exist.")
        print("  2. No readable PDF/TXT/DOCX files were found in the folder.")
        print("  3. There was an error extracting content from the files.")