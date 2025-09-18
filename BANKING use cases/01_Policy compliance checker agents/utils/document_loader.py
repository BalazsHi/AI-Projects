import logging
import os
from docx import Document

logger = logging.getLogger(__name__)

# Try to import PDF libraries with fallbacks
PDF_LIBRARY = None
try:
    import fitz  # PyMuPDF
    PDF_LIBRARY = "pymupdf"
    logger.info("Using PyMuPDF for PDF processing")
except ImportError:
    try:
        import pdfplumber
        PDF_LIBRARY = "pdfplumber"
        logger.info("Using pdfplumber for PDF processing")
    except ImportError:
        logger.error("No PDF library available. Install PyMuPDF or pdfplumber")

class DocumentLoader:
    """Utility class for loading various document formats."""
    
    @staticmethod
    def load_document(file_path: str) -> tuple[str, str]:
        """
        Load document content based on file extension.
        Returns: (content, error_message)
        """
        try:
            if not os.path.exists(file_path):
                error_msg = f"File not found: {file_path}"
                logger.error(error_msg)
                return "", error_msg
            
            file_size = os.path.getsize(file_path)
            logger.info(f"Loading document: {file_path} (Size: {file_size} bytes)")
            
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                content = DocumentLoader._load_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                content = DocumentLoader._load_docx(file_path)
            elif file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            # Debug: Print first 500 characters
            logger.info(f"Successfully loaded {len(content)} characters from {file_path}")
            logger.info(f"First 500 characters: {content[:500]}")
            
            if not content or len(content.strip()) < 10:
                error_msg = f"Document appears to be empty or too short: {len(content)} characters"
                logger.error(error_msg)
                return "", error_msg
                
            return content, ""
            
        except Exception as e:
            error_msg = f"Error loading document {file_path}: {str(e)}"
            logger.error(error_msg)
            return "", error_msg
    
    @staticmethod
    def _load_pdf(file_path: str) -> str:
        """Load text content from PDF file."""
        if PDF_LIBRARY == "pymupdf":
            return DocumentLoader._load_pdf_pymupdf(file_path)
        elif PDF_LIBRARY == "pdfplumber":
            return DocumentLoader._load_pdf_pdfplumber(file_path)
        else:
            raise ValueError("No PDF library available")
    
    @staticmethod
    def _load_pdf_pymupdf(file_path: str) -> str:
        """Load PDF using PyMuPDF."""
        import fitz
        
        doc = None
        try:
            logger.info(f"Opening PDF with PyMuPDF: {file_path}")
            doc = fitz.open(file_path)
            logger.info(f"PDF opened successfully. Pages: {len(doc)}")
            
            text_parts = []
            for page_num in range(len(doc)):
                try:
                    page = doc.load_page(page_num)  # Use load_page instead of indexing
                    page_text = page.get_text()
                    text_parts.append(f"\n--- Page {page_num + 1} ---\n{page_text}")
                    logger.info(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                except Exception as page_error:
                    logger.warning(f"Error reading page {page_num + 1}: {str(page_error)}")
                    continue
            
            text = "".join(text_parts)
            logger.info(f"Total PDF text extracted: {len(text)} characters")
            return text.strip()
            
        finally:
            if doc:
                try:
                    doc.close()
                    logger.info("PDF document closed successfully")
                except:
                    pass
    
    @staticmethod
    def _load_pdf_pdfplumber(file_path: str) -> str:
        """Load PDF using pdfplumber."""
        import pdfplumber
        
        try:
            logger.info(f"Opening PDF with pdfplumber: {file_path}")
            text_parts = []
            
            with pdfplumber.open(file_path) as pdf:
                logger.info(f"PDF opened successfully. Pages: {len(pdf.pages)}")
                
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text() or ""
                        text_parts.append(f"\n--- Page {page_num + 1} ---\n{page_text}")
                        logger.info(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                    except Exception as page_error:
                        logger.warning(f"Error reading page {page_num + 1}: {str(page_error)}")
                        continue
            
            text = "".join(text_parts)
            logger.info(f"Total PDF text extracted: {len(text)} characters")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error loading PDF with pdfplumber {file_path}: {str(e)}")
            raise e
    
    @staticmethod
    def _load_docx(file_path: str) -> str:
        """Load text content from DOCX file."""
        try:
            logger.info(f"Opening DOCX: {file_path}")
            doc = Document(file_path)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            logger.info(f"DOCX text extracted: {len(text)} characters")
            return text
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            raise e
