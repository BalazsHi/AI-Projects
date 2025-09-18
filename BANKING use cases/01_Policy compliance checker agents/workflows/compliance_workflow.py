from langgraph.graph import StateGraph, END
from typing import Dict, Any, List
from agents.supervisor import SupervisorAgent
from agents.chunker import ChunkerAgent
from agents.requirement_extractor import RequirementExtractorAgent
from agents.compliance_checker import ComplianceCheckerAgent
from utils.json_utils import JsonUtils
from utils.document_loader import DocumentLoader
from docx import Document
import os
import logging
from rich.console import Console
from rich.progress import Progress, TaskID

logger = logging.getLogger(__name__)
console = Console()

class ComplianceWorkflowState:
    def __init__(self):
        self.bank_policy = ""
        self.regulatory_document = ""
        self.chunks = []
        self.requirements = []
        self.compliance_results = []
        self.final_report = {}
        self.current_step = "initialization"
        self.progress = None

class ComplianceWorkflow:
    def __init__(self, api_key: str, output_dir: str = "output"):
        self.api_key = api_key
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
        # Initialize agents
        self.supervisor = SupervisorAgent(api_key)
        self.chunker = ChunkerAgent()
        self.requirement_extractor = RequirementExtractorAgent(api_key)
        self.compliance_checker = ComplianceCheckerAgent(api_key)
        
        self.state = ComplianceWorkflowState()
    
    def run_workflow(self, bank_policy_path: str, regulatory_doc_path: str) -> Dict[str, Any]:
        """Run the complete compliance workflow."""
        console.print("[bold green]Starting Compliance Workflow[/bold green]")
        
        with Progress() as progress:
            self.state.progress = progress
            main_task = progress.add_task("[green]Overall Progress", total=7)
            
            try:
                # Step 1: Load documents
                progress.update(main_task, description="[green]Loading documents...")
                self._load_documents(bank_policy_path, regulatory_doc_path)
                progress.update(main_task, advance=1)
                
                # Step 2: Assess document length
                progress.update(main_task, description="[green]Assessing document complexity...")
                assessment = self._assess_document_length()
                progress.update(main_task, advance=1)
                
                # Step 3: Chunk if necessary
                if assessment['needs_chunking']:
                    progress.update(main_task, description="[green]Chunking regulatory document...")
                    self._chunk_document(assessment['suggested_chunks'])
                else:
                    self._create_single_chunk()
                progress.update(main_task, advance=1)
                
                # Step 4: Extract requirements
                progress.update(main_task, description="[green]Extracting regulatory requirements...")
                self._extract_requirements()
                progress.update(main_task, advance=1)
                
                # Step 5: Check compliance
                progress.update(main_task, description="[green]Checking policy compliance...")
                self._check_compliance()
                progress.update(main_task, advance=1)
                
                # Step 6: Generate final report
                progress.update(main_task, description="[green]Generating final report...")
                self._generate_final_report()
                progress.update(main_task, advance=1)
                
                # Step 7: Save results
                progress.update(main_task, description="[green]Saving results...")
                self._save_results()
                progress.update(main_task, advance=1)
                
                console.print("[bold green]✓ Workflow completed successfully![/bold green]")
                return self.state.final_report
                
            except Exception as e:
                console.print(f"[bold red]✗ Workflow failed: {str(e)}[/bold red]")
                logger.error(f"Workflow failed: {str(e)}")
                raise
    
    def _load_documents(self, bank_policy_path: str, regulatory_doc_path: str):
        """Load both documents."""
        console.print("📄 Loading bank policy...")
        self.state.bank_policy, error = DocumentLoader.load_document(bank_policy_path)
        if error:
            raise Exception(f"Failed to load bank policy: {error}")
        
        console.print("📋 Loading regulatory document...")
        self.state.regulatory_document, error = DocumentLoader.load_document(regulatory_doc_path)
        if error:
            raise Exception(f"Failed to load regulatory document: {error}")
        
        console.print(f"✓ Loaded bank policy ({len(self.state.bank_policy)} chars)")
        console.print(f"✓ Loaded regulatory document ({len(self.state.regulatory_document)} chars)")
    
    def _assess_document_length(self) -> Dict[str, Any]:
        """Assess if document needs chunking."""
        console.print("🔍 Assessing document complexity...")
        assessment = self.supervisor.assess_document_length(self.state.regulatory_document)
        
        if assessment['needs_chunking']:
            console.print(f"📊 Document needs chunking into {assessment['suggested_chunks']} sections")
        else:
            console.print("📊 Document can be processed as single unit")
        
        return assessment
    
    def _chunk_document(self, num_chunks: int):
        """Chunk the regulatory document."""
        console.print(f"✂️ Chunking document into sections...")
    
        # Use the chunker's chunk_document method instead
        raw_chunks = self.chunker.chunk_document(
            self.state.regulatory_document, 
            doc_type="regulatory"
        )
    
        # Convert to the expected format
        self.state.chunks = []
        for chunk_data in raw_chunks:
            self.state.chunks.append({
                "chunk_id": chunk_data["chunk_id"],
                "content": chunk_data["content"],
                "start_position": 0,  # Could be calculated if needed
                "end_position": len(chunk_data["content"]),
                "section_info": f"Chunk {chunk_data['chunk_id']} of {len(raw_chunks)}"
            })
    
        console.print(f"✓ Created {len(self.state.chunks)} chunks")

    
    def _create_single_chunk(self):
        """Create single chunk for processing."""
        self.state.chunks = [{
            "chunk_id": 1,
            "content": self.state.regulatory_document,
            "start_position": 0,
            "end_position": len(self.state.regulatory_document),
            "section_info": "Complete Document"
        }]
    
    def _extract_requirements(self):
        """Extract requirements from all chunks."""
        console.print("🔍 Extracting regulatory requirements...")
    
        for i, chunk in enumerate(self.state.chunks):
            console.print(f"  Processing chunk {i+1}/{len(self.state.chunks)}")

            # Debug the chunk structure
            console.print(f"  Debug - Chunk keys: {list(chunk.keys())}")
        
            # Fix: Pass the entire chunk as a dictionary
            chunk_dict = {
                'content': chunk['content'],
                'section_info': chunk['section_info'],
                'chunk_id': chunk.get('chunk_id', f"chunk_{i+1}")
            }
            requirements_result = self.requirement_extractor.extract_requirements(chunk_dict)
        
            # Save intermediate JSON
            json_path = os.path.join(self.output_dir, f"requirements_chunk_{i+1}.json")
            JsonUtils.save_json(requirements_result, json_path)
        
            # Fix: Store the complete result (not append individual requirements)
            self.state.requirements.append(requirements_result)
    
        # Fix: Calculate total requirements properly
        total_reqs = 0
        for req_result in self.state.requirements:
            if isinstance(req_result, list):
                total_reqs += len(req_result)
            elif isinstance(req_result, dict) and 'requirements' in req_result:
                total_reqs += len(req_result['requirements'])
            else:
                total_reqs += 1  # Fallback count
    
        console.print(f"✓ Extracted {total_reqs} total requirements")

    
    def _check_compliance(self):
        """Check compliance for all requirements."""
        console.print("⚖️ Checking policy compliance...")
    
        for i, requirements_result in enumerate(self.state.requirements):
            console.print(f"  Checking compliance for chunk {i+1}/{len(self.state.requirements)}")
        
            # Debug: Print the structure
            console.print(f"  Debug - requirements_result type: {type(requirements_result)}")
            console.print(f"  Debug - requirements_result content: {str(requirements_result)[:200]}...")
        
            try:
                # Handle different result formats
                if isinstance(requirements_result, list):
                    requirements_to_check = requirements_result
                    console.print(f"  Debug - Using list format, length: {len(requirements_to_check)}")
                elif isinstance(requirements_result, dict) and 'requirements' in requirements_result:
                    requirements_to_check = requirements_result['requirements']
                    console.print(f"  Debug - Using dict format, requirements length: {len(requirements_to_check)}")
                else:
                    requirements_to_check = [requirements_result]  # Fallback
                    console.print(f"  Debug - Using fallback format")
            
                # Debug the structure we're passing to compliance checker
                console.print(f"  Debug - Passing to checker: type={type(requirements_to_check)}, length={len(requirements_to_check) if hasattr(requirements_to_check, '__len__') else 'unknown'}")
            
                compliance_result = self.compliance_checker.check_compliance(
                    self.state.bank_policy,
                    requirements_to_check
                )
                self.state.compliance_results.append(compliance_result)
            
            except Exception as e:
                console.print(f"  ✗ Error in chunk {i+1}: {str(e)}")
                logger.error(f"Error processing chunk {i+1}: {str(e)}")
                # Continue with next chunk instead of failing completely
                continue
    
        console.print("✓ Compliance check completed")


    
    def _generate_final_report(self):
        """Generate final compliance report."""
        console.print("📊 Generating final compliance report...")
        self.state.final_report = self.supervisor.create_final_report(
            self.state.compliance_results
        )
        console.print("✓ Final report generated")
    
    def _save_results(self):
        """Save all results to files."""
        console.print("💾 Saving results...")
        
        # Save JSON report
        json_path = os.path.join(self.output_dir, "compliance_report.json")
        JsonUtils.save_json(self.state.final_report, json_path)
        
        # Save Word document
        doc_path = os.path.join(self.output_dir, "compliance_report.docx")
        self._create_word_report(doc_path)
        
        console.print(f"✓ Results saved to {self.output_dir}/")
        console.print(f"  - JSON report: compliance_report.json")
        console.print(f"  - Word report: compliance_report.docx")
    
    def _create_word_report(self, file_path: str):
        """Create Word document report."""
        doc = Document()
        
        # Title
        title = doc.add_heading('Bank Policy Compliance Report', 0)
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(self.state.final_report.get('executive_summary', 'No summary available.'))
        
        # Statistics
        doc.add_heading('Compliance Statistics', level=1)
        breakdown = self.state.final_report.get('compliance_breakdown', {})
        for status, count in breakdown.items():
            doc.add_paragraph(f"{status.replace('_', ' ').title()}: {count}")
        
        # Detailed Findings
        doc.add_heading('Detailed Findings', level=1)
        for finding in self.state.final_report.get('detailed_findings', []):
            doc.add_heading(f"Requirement: {finding.get('id', 'N/A')}", level=2)
            doc.add_paragraph(f"Description: {finding.get('requirement', 'N/A')}")
            doc.add_paragraph(f"Reference: {finding.get('reference', 'N/A')}")
            doc.add_paragraph(f"Status: {finding.get('compliance_status', 'N/A').replace('_', ' ').title()}")
            
            if finding.get('recommendations'):
                doc.add_paragraph(f"Recommendations: {finding.get('recommendations', 'None')}")
        
        doc.save(file_path)
